"""DOCX 处理器"""
from pathlib import Path
from cleaner.state import upsert_metadata, log

BASE_DIR = Path.home() / "Documents" / "knowledge-system"


def process_docx(file_hash: str, vault_path: str, conn) -> dict:
    try:
        from docx import Document
    except ImportError:
        log(conn, file_hash, "processor", "error", "python-docx 未安装")
        return None

    doc = Document(vault_path)
    lines = []

    # 段落文本
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    # 表格文本
    for table in doc.tables:
        lines.append("")
        lines.append("| " + " | ".join("—" * 5 for _ in table.columns) + " |")
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    text = "\n".join(lines)
    text_path = BASE_DIR / "processed" / "text" / f"{file_hash}.txt"
    text_path.write_text(text, encoding="utf-8")

    # 内嵌图片
    image_paths = []
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    try:
        img_dir = BASE_DIR / "processed" / "images" / file_hash
        img_dir.mkdir(parents=True, exist_ok=True)
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.reltype:
                img = rel.target_part
                ext = Path(img.partname).suffix or ".png"
                img_path = img_dir / f"image_{i:03d}{ext}"
                img_path.write_bytes(img.blob)
                image_paths.append(str(img_path))
    except Exception:
        pass

    # 更新元数据
    props = doc.core_properties
    meta = {
        "title": props.title or Path(vault_path).stem,
        "author": props.author,
        "page_count": len(doc.paragraphs) // 30 or None,
    }
    upsert_metadata(conn, file_hash, **meta)

    log(conn, file_hash, "processor", "done", f"docx→text: {len(text)} chars, {len(image_paths)} images")
    return {"text_path": str(text_path), "image_paths": image_paths, "meta": meta}
